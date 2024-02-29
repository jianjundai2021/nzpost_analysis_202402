import os
from docx import Document


def add_paragraph_from_source(target_doc, source_para):
    target_para = target_doc.add_paragraph(style=source_para.style)
    for run in source_para.runs:
        target_run = target_para.add_run(run.text)
        # Copy formatting
        target_run.bold = run.bold
        target_run.italic = run.italic
        target_run.underline = run.underline
        target_run.font.color.rgb = run.font.color.rgb
        target_run.font.size = run.font.size
        # target_run.style = run.style
        if run.style:
            target_run.style = run.style 
		
def merge_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # If want to add page break. Don't add a page break if it's the first document
        # if index != 0:
        #     merged_document.add_page_break()

        for para in sub_doc.paragraphs:
            add_paragraph_from_source(merged_document, para)

        # add empty lines 
        if index < len(files) -1:
            merged_document.add_paragraph()
			
    return merged_document

folder_path = '.\working'  # Change this to your folder's path
files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith('.docx')]

# Sort files by name or any other criteria if necessary
files.sort()

merged_document = merge_word_documents(files)
merged_document.save('merged_document.docx')  # Save the merged document